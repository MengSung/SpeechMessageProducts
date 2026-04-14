from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPORT_NAME = "Server_Selection_Report_Lenovo_HPE_2026-04-12.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_doc_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    style.font.size = Pt(10.5)

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        s = doc.styles[style_name]
        s.font.name = "Calibri"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")

    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)


def set_run_font(run, name="Calibri", east_asia="Microsoft JhengHei", size=None, bold=None, italic=None):
    run.font.name = name
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), east_asia)


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Windows 伺服器採購分析報告\n")
    set_run_font(r, size=Pt(20), bold=True)
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    r2 = p.add_run("Lenovo / HPE 對標 Dell PowerEdge R660xs（Xeon Silver 4514Y、64GB / 128GB、2TB）")
    set_run_font(r2, size=Pt(12))

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p2.add_run(f"報告日期：{date(2026, 4, 12).isoformat()}")
    set_run_font(r3, italic=True)


def add_bullets(doc: Document, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers, rows, style="Table Grid"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = style
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], "D9EAF7")
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    return table


def add_source_paragraph(doc: Document, label: str, url: str):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(f"{label}: ").bold = True
    p.add_run(url)


def build_report():
    doc = Document()
    set_doc_defaults(doc)
    add_title(doc)

    doc.add_heading("一、先給結論", level=1)
    add_numbered(
        doc,
        [
            "最推薦採購：HPE ProLiant DL380 Gen11，建議配置為 1 顆 Xeon Silver 4514Y、128GB ECC DDR5、企業級 SSD 2x960GB 或 2x1.92TB、硬體 RAID、雙電源、3 年以上到府維保。",
            "如果重視規格與性價比：Lenovo ThinkSystem SR650 V3 是最強替代方案；若在地代理商與維保能力夠強，這台可以買。",
            "64GB 對四個產品屬於『可以啟用，但不適合作為長期穩定配置』；128GB 才是本案較保守、較可靠的容量。",
            "若以你提到的 Dell PowerEdge R660xs / Xeon Silver 4514Y 為對標，Lenovo 最接近的是 SR630 V3 / SR650 V3，HPE 最接近的是 DL360 Gen11 / DL380 Gen11；實際上 DL380 Gen11 與 SR650 V3 更適合四產品整合與未來擴充。",
            "『是否在中國製造』不能直接等於『品質較差』。Lenovo、HPE、Dell 都已經是全球供應鏈、多地製造與區域交付模式；真正影響長期可靠度的是平台成熟度、維保、零件供應、遠端管理、散熱與擴充規劃。"
        ],
    )

    doc.add_heading("二、需求定義與本報告判準", level=1)
    add_bullets(
        doc,
        [
            "作業系統：Windows Server 系列工作負載。",
            "目標配置：Xeon Silver 4514Y、64GB 或 128GB ECC DDR5、有效容量約 2TB。",
            "比較品牌：Lenovo 與 HPE，並以 Dell PowerEdge R660xs 作為對標參考。",
            "評估面向：CPU 能力、記憶體擴充性、儲存與 PCIe 擴充性、機箱密度、遠端管理、製造地與供應鏈、售後服務、價格與總體風險。",
            "時間尺度：本報告以 2026-04-12 可查得的公開規格與公開通路價格為基礎；企業實際成交價通常會因代理商折扣、保固年限、零件料號與地區而有顯著差異。"
        ],
    )

    doc.add_heading("三、核心 CPU 基準：Intel Xeon Silver 4514Y", level=1)
    cpu_table_rows = [
        ["處理器", "Intel Xeon Silver 4514Y"],
        ["世代", "第 5 代 Intel Xeon Scalable（Emerald Rapids）"],
        ["核心 / 執行緒", "16 核 / 32 執行緒"],
        ["基礎時脈 / Turbo", "2.0 GHz / 最高 3.4 GHz"],
        ["快取", "30 MB L3"],
        ["TDP", "150W"],
        ["記憶體支援", "DDR5，最多 8 通道 / CPU"],
        ["I/O", "PCIe 5.0，最多 80 lanes / CPU"],
        ["可擴充性", "2S（可單路起步、未來加第 2 顆 CPU）"],
        ["官方建議單價", "Intel RCP 約 US$780~790（僅 CPU，非整機）"],
    ]
    add_table(doc, ["項目", "內容"], cpu_table_rows)
    doc.add_paragraph(
        "判讀：4514Y 屬於中階、務實、低風險的企業 CPU，不是極高核心數型號，但對 Windows Server、資料庫、中小型虛擬化、ERP / Web / API / 報表等整合型工作負載相當合理。若你要的是穩定、功耗可控、後續可加第二顆 CPU，4514Y 是比高階大核心型號更平衡的選擇。"
    )

    doc.add_heading("四、64GB 與 128GB 的實務判斷", level=1)
    mem_rows = [
        [
            "64GB",
            "可跑輕到中度負載、少量 VM、少量 SQL / API / Web",
            "四產品若同機整合，Windows、資料庫、快取、背景作業同時發生時容易逼近瓶頸",
            "不建議作為長期正式配置；除非四產品都很輕且成長有限",
        ],
        [
            "128GB",
            "較適合四產品整合、保留緩衝、Windows Update、報表與尖峰查詢",
            "能降低記憶體壓力造成的磁碟換頁、效能抖動與擴容壓力",
            "本案建議值，也是我最推薦的起始容量",
        ],
    ]
    add_table(doc, ["容量", "適用情境", "主要風險 / 優勢", "結論"], mem_rows)
    doc.add_paragraph(
        "保守估計下，若四個產品中含 SQL Server、IIS / API、報表、背景排程、檔案交換或 3~4 台 VM，64GB 很容易只剩下『能運作』而不是『穩定舒服』。128GB 能明顯提高容錯與成長空間。"
    )

    doc.add_heading("五、品牌、製造地與『IBM 時代 Lenovo』的真實分析", level=1)
    add_bullets(
        doc,
        [
            "Lenovo 確實不是當年的 IBM System x，但不能直接推論現在品質一定比較差。現代 x86 伺服器是高度標準化供應鏈，可靠度更依賴平台設計、韌體成熟度、驗證流程、維保與零件供應。",
            "Lenovo 官方已公開其全球混合製造模式，包含中國、匈牙利、墨西哥、美國等多地；其匈牙利工廠明確用於伺服器與基礎架構產品。",
            "HPE 也不是單一國家製造。HPE 同樣採全球供應鏈與區域製造交付，並已公開捷克等歐洲製造投資。",
            "Dell 亦屬全球供應鏈模式，因此『非中國製造』不是 Lenovo 獨有優勢，也不是 HPE 或 Dell 的絕對優勢。",
            "真正應該比較的是：1) 當地代理商能力，2) 到府維修 SLA，3) 韌體與管理工具成熟度，4) 後續 3~7 年料件可得性，5) 機箱散熱與擴充是否能降低未來維護複雜度。"
        ],
    )
    doc.add_paragraph(
        "結論：若你要的是最保守、最接近『企業標準答案』的採購路線，HPE 的整體信心仍然略高；若你願意接受全球化製造現實、並有可靠代理商支撐，Lenovo 並不是品質不可買，反而在規格與擴充上很強。"
    )

    doc.add_heading("六、Lenovo / HPE / Dell 對標完整規格比較", level=1)
    compare_headers = [
        "品牌",
        "型號",
        "定位",
        "機箱",
        "CPU 擴充",
        "最大核心支援",
        "記憶體擴充",
        "儲存擴充",
        "PCIe / OCP",
        "遠端管理",
        "適合度",
    ]
    compare_rows = [
        [
            "Dell",
            "PowerEdge R660xs",
            "你的對標基準機",
            "1U / 2S",
            "最多 2 顆 Xeon",
            "依代次配置，高密度計算取向",
            "16 DIMM，容量較 2U 保守",
            "最多約 10x2.5 吋，1U 高密度",
            "約 3 PCIe + OCP",
            "iDRAC9",
            "適合機櫃密度優先；若要長期擴充，1U 比 2U 更吃規劃",
        ],
        [
            "Lenovo",
            "ThinkSystem SR630 V3",
            "最接近 R660xs 的 Lenovo 版",
            "1U / 2S",
            "最多 2 顆 4th/5th Gen Xeon",
            "每顆最高 64 核（依官方平台）",
            "32 DIMM，最高 8TB",
            "最多 10x2.5 吋或 16 EDSFF",
            "最多後部 3 PCIe + 前部 2 PCIe",
            "XClarity Controller 2",
            "若一定要 1U，這是 Lenovo 好選擇",
        ],
        [
            "Lenovo",
            "ThinkSystem SR650 V3",
            "Lenovo 的主力建議機",
            "2U / 2S",
            "最多 2 顆 4th/5th Gen Xeon",
            "每顆最高 64 核（依官方平台）",
            "32 DIMM，最高 8TB",
            "盤位彈性大，前中後擴充強",
            "最多約 10 PCIe + 1 OCP",
            "XClarity Controller 2",
            "四產品整合、未來升級、加卡加碟都比 1U 更友善",
        ],
        [
            "HPE",
            "ProLiant DL360 Gen11",
            "最接近 R660xs 的 HPE 版",
            "1U / 2S",
            "最多 2 顆 4th/5th Gen Xeon",
            "每顆最高 64 核",
            "每 CPU 16 DIMM，整機最高 8TB",
            "最高 20 EDSFF，配置靈活",
            "最多 2 x16 PCIe Gen5 + 2 OCP",
            "iLO 6",
            "1U 方案中非常成熟，但仍不如 DL380 好養",
        ],
        [
            "HPE",
            "ProLiant DL380 Gen11",
            "本報告總體首選",
            "2U / 2S",
            "最多 2 顆 4th/5th Gen Xeon",
            "每顆最高 64 核",
            "每 CPU 16 DIMM，整機最高 8TB",
            "8SFF / 24SFF / 8LFF / 12LFF 等多機型",
            "最多 8 PCIe Gen5 + 2 OCP",
            "iLO 6",
            "綜合可靠度、維保、擴充、可維護性最佳",
        ],
    ]
    add_table(doc, compare_headers, compare_rows)

    doc.add_heading("七、價格分析", level=1)
    doc.add_paragraph(
        "價格採用 2026-04-12 前後可查得之公開通路頁面或品牌商店資料。伺服器成交價會受保固、地區、代理商折扣、是否含導軌、RAID、網卡、電源與授權等條件影響，所以本表只能視為『公開市場參考價』，不可視為最終企業專案價。換算匯率以 1 USD ≈ 31.67 TWD 粗估。"
    )
    price_rows = [
        [
            "Lenovo SR650 V3",
            "1x Xeon Silver 4514Y / 64GB / 2x960GB SSD / RAID / 2x1100W",
            "US$6,060",
            "約 NT$191,900",
            "Microless 公開通路價，屬國際零售參考價",
        ],
        [
            "Lenovo SR650 V3",
            "2x Xeon 5415+ / 256GB",
            "US$16,309.99",
            "約 NT$516,500",
            "CDW 公開價，用於觀察高配後價格區間",
        ],
        [
            "HPE DL380 Gen11",
            "1x Xeon Silver 4514Y / 128GB / 2x480GB SSD / 2x1000W",
            "US$15,257",
            "約 NT$483,700",
            "CDW 公開價，通路與服務因素明顯墊高",
        ],
        [
            "HPE DL380 Gen11",
            "1x Xeon Silver 4410Y / 64GB / 2x960GB SSD",
            "US$4,981",
            "約 NT$157,700",
            "SHI 公開價，CPU 較低一級，用於看 HPE 入門價格帶",
        ],
        [
            "Dell R660xs",
            "官方起始配置",
            "US$3,249 起",
            "約 NT$102,900 起",
            "起跳價不等於你的 4514Y / 64GB / 2TB 實際配置價",
        ],
    ]
    add_table(doc, ["型號", "公開配置", "公開價格", "約當台幣", "解讀"], price_rows)
    doc.add_paragraph(
        "價格判讀：HPE 的公開價格波動通常比 Lenovo 更大，原因常包含服務包、企業級通路與 Smart Choice 組態；Lenovo 的街價常顯得更漂亮，但若把 128GB、企業 SSD、雙電源、正式 RAID、3~5 年到府保固全數補齊，總價差距會縮小。若你把停機成本與售後納入，HPE 的『貴』有一部分其實是買服務確定性。"
    )

    doc.add_heading("八、CPU 與未來擴充性重點", level=1)
    expand_rows = [
        ["起始方案", "1 顆 Xeon Silver 4514Y", "初期成本較低，適合中階工作負載"],
        ["未來 CPU 升級", "同平台可補第 2 顆 CPU", "需同時檢查散熱件、記憶體對稱與授權規劃"],
        ["記憶體升級", "64GB -> 128GB -> 256GB 都屬自然升級", "建議一開始就選較少條、較大容量的 ECC RDIMM，保留插槽"],
        ["儲存升級", "2x960GB 可升級到 2x1.92TB，或加資料碟 / 熱備援", "正式環境建議企業級 SSD / SAS，不建議消費級 SSD"],
        ["擴充卡", "2U 機型更適合 NIC、HBA、RAID、GPU 或備援卡", "2U 在散熱、可維護性與後續變更成本上更友善"],
    ]
    add_table(doc, ["項目", "建議", "說明"], expand_rows)

    doc.add_heading("九、售後服務與可靠度判讀", level=1)
    service_rows = [
        [
            "HPE",
            "iLO 6 生態成熟、企業市場滲透高、代理商與維保體系完整",
            "採購 3 年以上到府、最好 24x7 或 4 小時到場類型服務",
            "若你把停機風險放第一，HPE 仍是最保守答案",
        ],
        [
            "Lenovo",
            "新一代 Premier Support Plus for Infrastructure 強調 24x7、主動式與 AI 驅動支援",
            "高度取決於當地代理商實力與你實際購買的支援層級",
            "若代理商強、價格漂亮，SR650 V3 很值得買",
        ],
        [
            "Dell（對標參考）",
            "iDRAC 與 PowerEdge 平台也相當成熟",
            "本報告非主購品牌，但作為對標屬合理基準",
            "若只比 1U 密度，R660xs 很有競爭力",
        ],
    ]
    add_table(doc, ["品牌", "優勢", "購買時應要求", "判讀"], service_rows)
    doc.add_paragraph(
        "真正決定『可靠』的，不只是硬體本體，而是你買到哪一種保固與服務。若沒有至少 3 年到府維保、零件與工時明確 SLA，再好的品牌也可能在故障時讓你很痛苦。"
    )

    doc.add_heading("十、最終採購建議", level=1)
    final_rows = [
        [
            "A. 首選穩健方案",
            "HPE ProLiant DL380 Gen11",
            "1x Xeon Silver 4514Y、128GB ECC DDR5、2x960GB 或 2x1.92TB Enterprise SSD、硬體 RAID、雙電源、3 年以上到府保固",
            "最平衡、最保守、售後與維運風險最低",
        ],
        [
            "B. 高 CP 值方案",
            "Lenovo ThinkSystem SR650 V3",
            "1x Xeon Silver 4514Y、128GB ECC DDR5、2x960GB 或 2x1.92TB Enterprise SSD、RAID、雙電源、Premier 等級支援",
            "規格很強，價格通常更漂亮；前提是代理商與在地服務要可靠",
        ],
        [
            "C. 不建議長期方案",
            "任一品牌的 64GB 起始方案",
            "僅適合短期上線、負載很輕或明確半年內會加 RAM 的場景",
            "四產品同機整合下風險較高",
        ],
    ]
    add_table(doc, ["方案", "推薦機型", "建議配置", "結論"], final_rows)
    doc.add_paragraph(
        "本報告的最終答案很明確：如果你要『最可靠、最穩、最不容易後悔』，買 HPE ProLiant DL380 Gen11，而且直接上 128GB，不要停在 64GB。若你更想壓低預算且有可信任的 Lenovo 代理商，SR650 V3 才是 Lenovo 裡最值得買的型號。"
    )

    doc.add_heading("十一、建議詢價 BOM（可直接給廠商）", level=1)
    bom_rows = [
        [
            "HPE 建議 BOM",
            "DL380 Gen11 2U / 1x Xeon Silver 4514Y / 128GB ECC DDR5 / 2x960GB 或 2x1.92TB 企業 SSD / RAID / 雙電源 / 10GbE 或保留 OCP 升級 / 導軌 / 3 年 24x7 到府保固",
        ],
        [
            "Lenovo 建議 BOM",
            "SR650 V3 2U / 1x Xeon Silver 4514Y / 128GB ECC DDR5 / 2x960GB 或 2x1.92TB 企業 SSD / RAID / 雙電源 / XCC 高階管理授權視需求 / 導軌 / 3 年以上 Premier 支援",
        ],
        [
            "共同要求",
            "請廠商報出：料號、保固 SLA、零件到場時間、是否含上架導軌、是否含 RAID 快取、SSD 寫入耐久度、第二顆 CPU 未來升級成本、128GB 升 256GB 的價格、同規格 5 年保固差額。",
        ],
    ]
    add_table(doc, ["項目", "內容"], bom_rows)

    doc.add_heading("十二、資料來源", level=1)
    sources = [
        ("Intel Xeon Silver 4514Y 官方規格", "https://www.intel.com/content/www/us/en/products/sku/237557/intel-xeon-silver-4514y-processor-30m-cache-2-00-ghz/specifications.html"),
        ("Lenovo ThinkSystem SR630 V3 官方規格", "https://pubs.lenovo.com/sr630-v3/server_specifications_technical"),
        ("Lenovo ThinkSystem SR650 V3 官方規格", "https://pubs.lenovo.com/sr650-v3/server_specifications_technical"),
        ("HPE ProLiant DL360 Gen11 官方頁", "https://www.hpe.com/us/en/products/compute/hpe-proliant-compute/dl360-gen11.html"),
        ("HPE ProLiant DL380 Gen11 官方頁", "https://www.hpe.com/us/en/products/compute/hpe-proliant-compute/dl380-gen11.html"),
        ("Dell PowerEdge R660xs 官方頁", "https://www.dell.com/en-us/shop/productdetailstxn/poweredge-r660xs"),
        ("Lenovo Premier Support Plus for Infrastructure", "https://www.lenovo.com/us/en/services/support-services/premier-support-plus-for-infrastructure/"),
        ("Lenovo 新一代伺服器 Premier Support Plus 新聞稿（2026-02-26）", "https://news.lenovo.com/pressroom/press-releases/always-on-infrastructure-premier-support-plus-proactive-ai/"),
        ("Lenovo 匈牙利伺服器製造工廠新聞稿", "https://news.lenovo.com/pressroom/press-releases/first-european-in-house-manufacturing-facility-ullo-hungary/"),
        ("HPE 捷克製造投資新聞稿", "https://www.hpe.com/us/en/newsroom/press-release/2022/05/hewlett-packard-enterprise-strengthens-europes-supercomputer-supply-chain-with-new-factory-in-czech-republic.html"),
        ("HPE DL380 Gen11 4514Y 128GB CDW 公開價格頁", "https://www.cdw.com/product/hpe-proliant-dl380-gen11-network-choice-rack-mountable-xeon-silver-4514/7931767"),
        ("Lenovo SR650 V3 4514Y 64GB 公開價格頁", "https://global.microless.com/product/lenovo-thinksystem-sr650-v3-2u-rack-server-intel-xeon-silver-4514y-64gb-ram-2x-960gb-dc600m-sata-ssd-raid-9350-8i-2gb-flash-pcie-12gb-adapter-2x-1100w-sr650-v3/?currency=usd"),
        ("SHI HPE DL380 Gen11 64GB / 2x960GB 參考頁", "https://www.shi.com/product/50381552/HPE-DL380-G11-4410Y-2X32G-8SFF-SSD-SVR"),
        ("美元兌台幣歷史匯率參考", "https://www.currency-converter.org.uk/currency-rates/historical/rate/USD-TWD-04_02_2026.html"),
    ]
    for label, url in sources:
        add_source_paragraph(doc, label, url)

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.add_run("附註：").bold = True
    note.add_run(
        "若要取得最終可下單價格，仍建議同步向台灣 Lenovo 與 HPE 授權代理商各取 1 份 128GB / 2TB / 3~5 年維保的正式報價單，再比對保固 SLA、SSD 型號、RAID 規格與第二顆 CPU 升級成本。"
    )

    out_path = Path.cwd() / REPORT_NAME
    doc.save(out_path)
    print(out_path)


if __name__ == "__main__":
    build_report()
