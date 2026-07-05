from __future__ import annotations

from datetime import date
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[3]
OUT_DIR = ROOT / "ChurchReport" / "文件" / "功能說明"
OUT_PATH = OUT_DIR / "LINE_RichMenu_創意與程式調用完整說明.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color: str = "D9E2F3", size: str = "4") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths: Iterable[float]) -> None:
    widths = list(widths)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def set_paragraph_spacing(paragraph, before: int = 0, after: int = 6, line: int = 300) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line / 240


def set_font(run, name: str = "Calibri", size: int | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, before=0, after=3, line=300)
    r = p.add_run("LINE RichMenu 創意、程式修改與調用完整說明")
    set_font(r, size=24, bold=True, color="1F4D78")

    subtitle = doc.add_paragraph()
    set_paragraph_spacing(subtitle, after=10, line=300)
    r = subtitle.add_run(f"適用分支：Jesus_5.1.7.RefactorRichMenu｜產出日期：{date.today().isoformat()}｜專案：ChurchReport")
    set_font(r, size=10, color="555555")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    r = p.add_run(text)
    set_font(r, size={1: 16, 2: 13, 3: 12}.get(level, 11), bold=True, color="2E74B5" if level < 3 else "1F4D78")
    set_paragraph_spacing(p, before={1: 18, 2: 12, 3: 8}.get(level, 6), after={1: 8, 2: 6, 3: 4}.get(level, 4), line=300)


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=6, line=300)
    r = p.add_run(text)
    set_font(r, size=11, color="000000")


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_paragraph_spacing(p, after=4, line=300)
        r = p.add_run(item)
        set_font(r, size=11, color="000000")


def add_numbered(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        set_paragraph_spacing(p, after=4, line=300)
        r = p.add_run(item)
        set_font(r, size=11, color="000000")


def add_code(doc: Document, code: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.3])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F9FC")
    set_cell_borders(cell, "DADCE0", "4")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=0, line=280)
    r = p.add_run(code.strip())
    set_font(r, name="Consolas", size=9, color="1F1F1F")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, "E8EEF5")
        set_cell_borders(cell)
        p = cell.paragraphs[0]
        r = p.add_run(header)
        set_font(r, size=10, bold=True, color="1F4D78")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_borders(cells[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cells[idx].paragraphs[0]
            set_paragraph_spacing(p, after=2, line=280)
            r = p.add_run(value)
            set_font(r, size=9, color="000000")
    doc.add_paragraph()


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    normal.font.size = Pt(11)


def build_doc() -> Document:
    doc = Document()
    configure_document(doc)
    add_title(doc)

    add_heading(doc, "一、文件目的與閱讀方式")
    add_body(
        doc,
        "這份文件把三件事放在同一份 Word 裡：第一，整理 LINE RichMenu 官方能力與可延伸的新創意；第二，說明本分支對 RichMenu 程式架構做了哪些修改；第三，提供開發者與系統維運者可以直接參考的調用方式。"
    )
    add_body(
        doc,
        "為避免誤解，文件會把狀態分成三類：已完成並已有程式入口、底層 SDK 已支援但共用 RichMenu workflow 尚未封裝、以及未來可按目前架構延伸的產品創意。"
    )

    add_heading(doc, "二、網路與官方資料整理")
    add_body(
        doc,
        "本次參考 LINE Developers 官方 RichMenu 文件、Messaging API reference、rich menu switch action 文件，以及 LIFF overview。官方能力重點包括：RichMenu 是 LINE 聊天室底部的可點擊選單；可設定多個點擊區域；可用 URI、message、postback、datetime picker、camera、camera roll、location、rich menu switch 等 action；也可透過 alias 讓選單內按鈕切換到另一個 RichMenu。"
    )
    add_table(
        doc,
        ["資料來源", "本文採用重點", "對目前程式的意義"],
        [
            ["LINE Developers - Use rich menus", "RichMenu 的版面、圖片、點擊區域、建立與使用流程。", "支撐 LineRichMenuDefinition 與 Provisioning workflow 的設計。"],
            ["LINE Messaging API reference - Rich menu", "Create、upload image、link/unlink user、default rich menu、alias、batch 等 API。", "支撐 LineMessagingProcessorClass 與 Line.Messaging SDK 的 API 封裝。"],
            ["LINE Messaging API reference - richmenu switch action", "透過 richMenuAliasId 與 postback data 讓使用者在 LINE App 內切換 RichMenu。", "支撐 RichMenuActionFactory.SwitchToAlias 與未來分頁式選單。"],
            ["LINE Developers - LIFF overview", "LIFF 可在 LINE 內開啟 Web app，適合從 RichMenu 入口導向會員中心、表單、查詢、上傳等互動頁。", "支撐 RichMenu 作為入口，LIFF 作為深層互動頁的產品設計。"],
        ],
        [1.65, 2.55, 2.1],
    )

    add_heading(doc, "三、LINE RichMenu 可以做什麼")
    add_table(
        doc,
        ["功能類型", "說明", "目前專案狀態"],
        [
            ["基本選單", "在聊天室下方顯示圖片式選單，使用者點擊不同區域觸發 action。", "已由 Line.Messaging RichMenu 型別與 LineRichMenuDefinition 描述。"],
            ["建立與上傳圖片", "先建立 RichMenu metadata/layout，再上傳 PNG/JPEG 圖片。", "已由 LineRichMenuProvisioningWorkflow 執行 create/upload。"],
            ["個人指派/解除指派", "將指定 RichMenu 綁定到特定 LINE user，或解除後回到 default menu。", "已由 ILineRichMenuAssignmentWorkflow 封裝，ChurchReport 已接入 Add/DeleteRichMenuMessage。"],
            ["預設選單", "設定 channel 層級的 default RichMenu。", "Provisioning workflow 可依 IsDefault 設定。"],
            ["Alias 與分頁切換", "用 alias 指向 RichMenu，並用 rich menu switch action 在選單間切換。", "alias 已在 provisioning 維護；RichMenuActionFactory 已提供 switch action helper。"],
            ["文字觸發切換", "收到特定文字後切換到對應 menuKey。", "LineRichMenuTextTriggerResolver 與 Policy 已完成，但 ChurchReport 尚未把 webhook 文字事件接入 Orchestrator。"],
            ["依角色/狀態決策", "依會員、同工、付款狀態、活動階段等 policy 決定選單。", "IRichMenuPolicy/IRichMenuOrchestrator 已完成，產品端需新增自己的 policy。"],
            ["暫時選單與到期恢復", "臨時活動、付款提醒、維修流程結束後恢復上一個選單。", "IRichMenuStateStore 與 RichMenuExpirationSweepWorkflow 已完成，需排程接入。"],
            ["批次指派/批次替換", "一次處理多位 user 或 batch replace。", "Line.Messaging SDK 已有 API/model；RichMenus 共用 workflow 尚未封裝。"],
            ["部署前驗證", "建立前先 validate rich menu 或 batch request。", "Line.Messaging SDK 已有 validate API；RichMenus provisioning 尚未在流程中自動呼叫。"],
        ],
        [1.25, 2.7, 2.35],
    )

    add_heading(doc, "四、網路研究整理出的新創意點子")
    add_body(doc, "以下點子不是全部都已在 ChurchReport 完成，而是依 LINE 官方能力、目前共用架構與 ChurchReport 場景整理出的可落地設計方向。")
    add_table(
        doc,
        ["創意點子", "適用場景", "可用目前程式如何落地", "狀態"],
        [
            ["分頁式 RichMenu", "主選單、會員資料、奉獻/繳費、課程/活動分頁。", "用 Alias + RichMenuActionFactory.SwitchToAlias 產生 rich menu switch action，postback data 採 richmenu:tab:from->to。", "alias 與 action helper 已有，需設計多張 catalog menu。"],
            ["身份/角色選單", "會友、同工、小組長、管理者看到不同入口。", "新增 IRichMenuPolicy，依 CRM/會員資料產生 RichMenuDecision.Assign(menuKey)。", "Orchestrator 已有，產品 policy 待實作。"],
            ["關鍵字切換選單", "使用者輸入「奉獻」、「課程」、「會員中心」即切換選單。", "AddLineRichMenus 設定 ExactTextToMenuKey，再由 webhook 建立 RichMenuContext 丟給 ApplyAsync。", "Resolver/Policy 已有，webhook 接入待實作。"],
            ["暫時活動選單", "營會、特會、報名期間、公告期間顯示活動入口。", "RichMenuDecision.Assign(menuKey, ttl: ...)，狀態寫入 IRichMenuStateStore，到期由 SweepAsync 恢復。", "State/Sweep 已有，TTL policy 與排程待接入。"],
            ["付款狀態選單", "未付款顯示繳費入口，付款後切到收據/報名完成入口。", "付款 callback 後呼叫 AssignOrThrowAsync(lineUserId, \"payment-complete\")。", "Assignment 已可用，需新增 catalog/menuKey。"],
            ["LIFF 深層互動入口", "會員資料維護、大頭照上傳、報名、奉獻查詢。", "RichMenu 區域使用 URI action 指向 LIFF URL；LIFF 做表單與權限流程。", "LINE 能力已支持，需產品頁與 URL 設計。"],
            ["公告/緊急通知選單", "活動當週、颱風公告、場地變更、特殊提醒。", "新增 high-priority policy 依活動狀態覆蓋一般 menu；活動結束後 sweep 恢復。", "架構支持，產品 policy 待實作。"],
            ["A/B 測試選單", "比較不同 CTA、圖片、分頁文字對點擊與轉換影響。", "用不同 menuKey/alias 分群，透過 policy 或批次指派分配。", "SDK 與架構可支持，分析追蹤需另做。"],
            ["批次族群切換", "課程名單、活動名單、同工名單一次切換到指定 RichMenu。", "底層 Line.Messaging 已有 bulk/batch API；共用 RichMenus 可新增 AssignManyAsync。", "目前屬可擴充能力。"],
            ["防呆部署流程", "RichMenu 上線前先檢查圖片、版面、action、alias 與 default。", "在 provisioning 前導入 ValidateRichMenuAsync/ValidateRichMenuBatchRequestAsync。", "底層 SDK 已有，workflow 待封裝。"],
        ],
        [1.3, 1.35, 2.6, 1.05],
    )

    add_heading(doc, "五、本分支 RichMenu 程式到底修改了什麼")
    add_heading(doc, "5.1 新增 LineMessagingProcessor.RichMenus 共用專案", 2)
    add_body(
        doc,
        "新增 LineMessagingProcessor.RichMenus 專案，目標是把 RichMenu 的共用流程從 ChurchReport 產品程式抽離出來。這個專案不應依賴 ChurchReport、CRM、Controller、DbContext 或特定產品語意，而是提供 menu catalog、佈建、指派、決策與狀態管理等通用能力。"
    )
    add_table(
        doc,
        ["元件", "責任", "主要檔案"],
        [
            ["Catalog", "產品宣告有哪些 RichMenu、alias、圖片 stream、是否 default。", "ILineRichMenuCatalog.cs、LineRichMenuDefinition.cs、StaticLineRichMenuCatalog.cs"],
            ["Provisioning", "讀取 catalog，依 fingerprint 判斷是否已存在，建立/上傳/alias/default/cache。", "ILineRichMenuProvisioningWorkflow.cs、LineRichMenuProvisioningWorkflow.cs、LineRichMenuFingerprint.cs"],
            ["Assignment", "把 menuKey 解析成 richMenuId，呼叫 LINE link/unlink，並記錄目前/前一個選單狀態。", "ILineRichMenuAssignmentWorkflow.cs、LineRichMenuAssignmentWorkflow.cs、LineRichMenuAssignmentResult.cs"],
            ["Orchestrator/Policy", "讓多個產品規則競爭決定應套用哪個 RichMenu。", "IRichMenuOrchestrator.cs、RichMenuOrchestrator.cs、IRichMenuPolicy.cs、RichMenuDecision.cs"],
            ["Text Trigger", "把使用者文字訊息對應到 menuKey。", "LineRichMenuTextTriggerOptions.cs、LineRichMenuTextTriggerResolver.cs、LineRichMenuTextTriggerPolicy.cs"],
            ["State/Sweep", "保存暫時選單狀態與到期時間，到期後恢復前一個選單或解除綁定。", "IRichMenuStateStore.cs、InMemoryRichMenuStateStore.cs、RichMenuExpirationSweepWorkflow.cs"],
            ["Action Factory", "產生 rich menu switch action，避免產品端手寫 SDK 細節。", "RichMenuActionFactory.cs"],
        ],
        [1.45, 3.0, 1.85],
    )

    add_heading(doc, "5.2 LineMessagingProcessor 與 Line.Messaging 的修改", 2)
    add_bullets(
        doc,
        [
            "LineMessagingProcessorClass 新增/整理 RichMenu 相關 API 封裝，例如 CreateRichMenuAsync、UploadRichMenuPngImageAsync、GetRichMenuListAsync、SetDefaultRichMenuAsync、GetDefaultRichMenuIdAsync、CancelDefaultRichMenuAsync、alias CRUD、link/unlink user。",
            "LineMessagingProcessorRichMenuAdapter 把 LineMessagingProcessorClass 轉接成 ILineRichMenuProcessor，讓 RichMenus 共用層不直接依賴具體 processor 實作。",
            "Line.Messaging 已具備 RichMenuSwitchTemplateAction、RichMenuBulkLinkRequest、RichMenuBatchOperation、RichMenuBatchProgress 等 SDK 型別與 API。",
            "需要特別標示：bulk/batch/validate 雖然在 Line.Messaging SDK 層存在，但目前 LineMessagingProcessor.RichMenus 共用 workflow 還沒有把它們包成 AssignMany、ReplaceLinkedMenu 或 ValidateBeforeCreate 這類高階流程。"
        ],
    )

    add_heading(doc, "5.2.1 原始碼註解與編碼注意事項", 3)
    add_body(
        doc,
        "本次 RichMenu 相關程式的註解內容是繁體中文，正確解碼方式是 UTF-8。若在 PowerShell、舊版命令提示字元或某些 log viewer 中看到亂碼，通常是終端機 code page 或輸出編碼設定問題，不代表原始碼或 Word 文件內容不是繁體中文。閱讀或轉檔時建議使用 UTF-8，必要時先切到 chcp 65001 或使用支援 UTF-8 的編輯器。"
    )

    add_heading(doc, "5.3 ASP.NET Core DI 修改", 2)
    add_body(
        doc,
        "LineMessagingProcessor.AspNetCore 新增 AddLineRichMenus 與 AddLineRichMenuProvisioning<TCatalog>。AddLineMessagingProcessor 會自動加入產品中立 RichMenu 服務；如果產品要讓 catalog 參與佈建，則另外呼叫 AddLineRichMenuProvisioning<TCatalog>。"
    )
    add_code(
        doc,
        """
services.AddLineMessagingProcessor(options =>
{
    options.ChannelAccessToken = configuration["LineMessaging:Jesus:ChannelAccessToken"];
    options.ApiBaseUri = "https://api.line.me/v2";
});

// 有 catalog 並需要同步到 LINE 後台時才註冊 provisioning。
services.AddLineRichMenuProvisioning<ChurchReportLegacyRichMenuCatalog>();

// 可選：設定文字觸發。
services.AddLineRichMenus(options =>
{
    options.ExactTextToMenuKey["會員中心"] = "member-main";
    options.ExactTextToMenuKey["奉獻"] = "donation-main";
});
"""
    )

    add_heading(doc, "5.4 ChurchReport 整合修改", 2)
    add_bullets(
        doc,
        [
            "新增 ChurchReportLegacyRichMenuCatalog，將既有單鈕認證 RichMenu 宣告為 menuKey = legacy-auth 的 catalog 項目。",
            "PushUtility.AddRichMenuMessage 不再自己建立/上傳/連結 RichMenu，而是呼叫 ILineRichMenuAssignmentWorkflow.AssignOrThrowAsync(UserId, \"legacy-auth\")。",
            "PushUtility.DeleteRichMenuMessage 改為呼叫 ILineRichMenuAssignmentWorkflow.UnassignOrThrowAsync(UserId)。",
            "LineUtilityClass.AddRichMenuMessage / DeleteRichMenuMessage 也改走相同 assignment workflow。",
            "這代表 ChurchReport 目前已接入的是「指派/解除指派」能力；Provisioning、Orchestrator、Text Trigger、Sweep 還需要由 controller、job 或 webhook 流程實際呼叫。"
        ],
    )

    add_heading(doc, "六、怎麼調用/呼叫這些 RichMenu 能力")
    add_heading(doc, "6.1 第一次部署或選單改版：同步 RichMenu 到 LINE", 2)
    add_body(doc, "前提是已註冊 AddLineRichMenuProvisioning<TCatalog>，且 catalog 的圖片路徑/stream 可讀。同步後會得到每個 menuKey 對應的 richMenuId，並寫入 cache。")
    add_code(
        doc,
        """
public sealed class RichMenuSyncJob
{
    private readonly ILineRichMenuProvisioningWorkflow _workflow;

    public RichMenuSyncJob(ILineRichMenuProvisioningWorkflow workflow)
    {
        _workflow = workflow;
    }

    public async Task RunAsync(CancellationToken cancellationToken = default)
    {
        LineRichMenuSyncReport report = await _workflow.SyncAsync(cancellationToken);

        foreach (var item in report.Items)
        {
            // item.Outcome: Created / UpToDate / Failed
            // item.MenuKey: 例如 legacy-auth
            // item.RichMenuId: LINE 平台 richMenuId
        }
    }
}
"""
    )

    add_heading(doc, "6.2 指派使用者到指定 RichMenu", 2)
    add_body(doc, "這是 ChurchReport 目前已經接入的核心路徑。若 menuKey 尚未同步或無法在線上找到對應版本，AssignAsync 會回傳 ValidationFailed；AssignOrThrowAsync 會丟 LineRichMenuException。")
    add_code(
        doc,
        """
await assignmentWorkflow.AssignOrThrowAsync(lineUserId, "legacy-auth");

LineRichMenuAssignmentResult result =
    await assignmentWorkflow.AssignAsync(lineUserId, "member-main");

if (!result.Succeeded)
{
    // result.Status: ValidationFailed / ProviderRejected / ProviderUnavailable / UnexpectedError
    // result.ErrorCode 與 ErrorMessage 可寫入 log 或回報管理者。
}
"""
    )

    add_heading(doc, "6.3 解除使用者 RichMenu", 2)
    add_body(doc, "解除個人 RichMenu 後，LINE 會回到 channel default RichMenu；如果沒有 default，則不顯示 RichMenu。")
    add_code(
        doc,
        """
await assignmentWorkflow.UnassignOrThrowAsync(lineUserId);

LineRichMenuAssignmentResult result =
    await assignmentWorkflow.UnassignAsync(lineUserId);
"""
    )

    add_heading(doc, "6.4 ChurchReport 目前既有呼叫方式", 2)
    add_body(doc, "如果維持原本 ChurchReport 工具類使用方式，呼叫點不需要知道 richMenuId，只要提供 LINE user id。內部會轉到共用 assignment workflow。")
    add_code(
        doc,
        """
var pushUtility = new PushUtility(lineMessagingClient);

// 指派 legacy-auth RichMenu。
await pushUtility.AddRichMenuMessage(lineUserId);

// 解除個人 RichMenu。
await pushUtility.DeleteRichMenuMessage(lineUserId);

// LineUtilityClass 也保留相同方法名稱：
await lineUtilityClass.AddRichMenuMessage(lineUserId);
await lineUtilityClass.DeleteRichMenuMessage(lineUserId);
"""
    )

    add_heading(doc, "6.5 文字觸發與 Orchestrator", 2)
    add_body(doc, "目前共用層已可把收到的文字對應到 menuKey；產品端需要在 webhook 收到文字訊息時建立 RichMenuContext 並呼叫 ApplyAsync。")
    add_code(
        doc,
        """
var context = new RichMenuContext(
    lineUserId: lineUserId,
    receivedText: messageText,
    currentMenuKey: currentMenuKey);

LineRichMenuAssignmentResult result =
    await richMenuOrchestrator.ApplyAsync(context, cancellationToken);
"""
    )

    add_heading(doc, "6.6 Alias 分頁切換 Action", 2)
    add_body(doc, "在 RichMenu 的某個 action area 中放入 rich menu switch action，即可讓使用者在 LINE App 中切換到另一張以 alias 管理的 RichMenu。")
    add_code(
        doc,
        """
var action = RichMenuActionFactory.SwitchToAlias(
    aliasId: "member-settings",
    data: "richmenu:tab:member-main->member-settings",
    label: "設定");
"""
    )

    add_heading(doc, "6.7 暫時選單到期恢復", 2)
    add_body(doc, "若 policy 指派暫時選單時有寫入 ExpiresAt，排程可定期呼叫 SweepAsync。到期後 workflow 會嘗試恢復 PreviousMenuKey；若沒有上一個選單，則解除個人 RichMenu。")
    add_code(
        doc,
        """
public sealed class RichMenuSweepJob
{
    private readonly IRichMenuExpirationSweepWorkflow _sweepWorkflow;

    public RichMenuSweepJob(IRichMenuExpirationSweepWorkflow sweepWorkflow)
    {
        _sweepWorkflow = sweepWorkflow;
    }

    public Task<RichMenuExpirationSweepReport> RunAsync()
    {
        return _sweepWorkflow.SweepAsync(DateTimeOffset.UtcNow);
    }
}
"""
    )

    add_heading(doc, "七、目前完成度與界線")
    add_table(
        doc,
        ["能力", "目前完成度", "注意事項"],
        [
            ["Create/Upload/List/Link/Unlink", "共用層已可使用。", "透過 ILineRichMenuProcessor 與 workflow 間接呼叫。"],
            ["Alias 維護", "Provisioning 已使用 create/update alias。", "用 alias 支撐分頁切換與穩定指向。"],
            ["Default RichMenu", "Provisioning 可依 IsDefault 設定。", "解除個人 RichMenu 後會回到 default。"],
            ["Assignment", "已完成且 ChurchReport 已接入。", "AddRichMenuMessage/DeleteRichMenuMessage 已轉到共用 workflow。"],
            ["Orchestrator/Policy", "共用層已完成。", "ChurchReport 尚未新增具體角色/狀態 policy 接入點。"],
            ["Text Trigger", "Resolver/Policy 已完成。", "Webhook 文字事件尚需呼叫 ApplyAsync。"],
            ["Expiration Sweep", "共用層已完成。", "需產品端排程或背景服務接入。"],
            ["Bulk/Batch/Validate", "Line.Messaging SDK 已支援。", "RichMenus 高階 workflow 尚未封裝，文件不可說成已完成。"],
            ["持久化狀態", "目前預設 InMemory。", "正式環境若需跨機/重啟保留，應改實作 Redis/DB 版 IRichMenuStateStore。"],
            ["ChurchReportLegacyRichMenuCatalog 圖片", "已新增 catalog。", "需確認部署環境可讀圖片檔，否則 SyncAsync 會失敗。"],
        ],
        [1.65, 2.0, 2.65],
    )

    add_heading(doc, "八、錯誤處理與維運提醒")
    add_bullets(
        doc,
        [
            "ValidationFailed 通常代表輸入不合法、menuKey 不存在、尚未 provisioning、或 catalog/圖片資料不足。",
            "ProviderRejected 代表 LINE API 有回應但拒絕，例如參數錯、token 權限不足、richMenuId 無效。",
            "ProviderUnavailable 代表 HTTP/網路/timeout 類問題，通常可重試或告警。",
            "UnexpectedError 不應吞掉，代表程式資料流或未知錯誤，需回頭修正程式。",
            "正式部署前應先跑 SyncAsync 並檢查 report.Items，確認所有必要 menuKey 都是 Created 或 UpToDate。",
            "若使用 InMemory cache/state store，服務重啟後 assignment 可能需要重新查 LINE RichMenu list 或重新同步。",
            "RichMenu alias 可降低 richMenuId 改版帶來的切換問題，但 alias 本身仍要由 provisioning 正確建立與更新。",
        ],
    )

    add_heading(doc, "九、建議導入順序")
    add_numbered(
        doc,
        [
            "先確認 ChurchReportLegacyRichMenuCatalog 的 PNG 圖片在部署環境可讀。",
            "在系統啟動後或管理者工具中呼叫 ILineRichMenuProvisioningWorkflow.SyncAsync。",
            "確認 report.Items 沒有 Failed，並保存 menuKey/richMenuId 對照到 log。",
            "保留既有 AddRichMenuMessage/DeleteRichMenuMessage 呼叫點，讓它們透過 assignment workflow 操作 legacy-auth。",
            "新增第二張或第三張 RichMenu 時，先擴充 catalog，再加入 alias 與 RichMenuActionFactory switch action。",
            "要做文字切換時，先設定 ExactTextToMenuKey，再在 webhook text message 路徑呼叫 IRichMenuOrchestrator.ApplyAsync。",
            "要做暫時活動選單時，新增 policy 與到期時間，並用背景排程呼叫 SweepAsync。",
            "要做大量名單切換時，再把 Line.Messaging 既有 bulk/batch API 包成 RichMenus 共用層高階方法。",
        ],
    )

    add_heading(doc, "十、測試與驗證現況")
    add_body(
        doc,
        "本分支已有 RichMenu 相關測試，涵蓋 provisioning 建立/重用/失敗續跑、assignment cache/online resolution/錯誤轉換、orchestrator 無變更與 policy 決策、text trigger trim + exact mapping、ChurchReport PushUtility 透過 assignment workflow 指派 legacy-auth 等案例。"
    )
    add_body(
        doc,
        "文件層面的重點不是重新證明所有程式測試，而是讓維運者知道：同步、指派、解除、分頁切換、文字觸發、到期恢復各自該呼叫哪個介面，以及哪些能力目前只是 SDK 支援但尚未進入共用 workflow。"
    )

    add_heading(doc, "十一、附錄：主要檔案索引")
    add_table(
        doc,
        ["檔案", "用途"],
        [
            ["LineMessagingProcessor.RichMenus/LineRichMenuProvisioningWorkflow.cs", "同步 catalog 到 LINE，建立 RichMenu、上傳圖片、維護 alias/default/cache。"],
            ["LineMessagingProcessor.RichMenus/LineRichMenuAssignmentWorkflow.cs", "指派/解除個人 RichMenu，解析 menuKey，標準化 LINE provider 錯誤。"],
            ["LineMessagingProcessor.RichMenus/RichMenuOrchestrator.cs", "整合多個 IRichMenuPolicy，選出最高優先權決策並交給 assignment workflow。"],
            ["LineMessagingProcessor.RichMenus/LineRichMenuTextTriggerResolver.cs", "依 ExactTextToMenuKey 做文字到 menuKey 的解析。"],
            ["LineMessagingProcessor.RichMenus/RichMenuExpirationSweepWorkflow.cs", "掃描過期 RichMenu 狀態並恢復上一個選單或解除綁定。"],
            ["LineMessagingProcessor.AspNetCore/LineMessagingProcessorServiceCollectionExtensions.cs", "註冊 AddLineRichMenus 與 AddLineRichMenuProvisioning<TCatalog>。"],
            ["ChurchReport/Tools/ChurchReportLegacyRichMenuCatalog.cs", "宣告 ChurchReport 既有 legacy-auth RichMenu。"],
            ["ChurchReport/Tools/PushUtility.cs", "AddRichMenuMessage/DeleteRichMenuMessage 改走共用 assignment workflow。"],
            ["ChurchReport/Tools/LineUtilityClass.cs", "保留既有方法名，但 RichMenu 指派改走共用 assignment workflow。"],
        ],
        [2.95, 3.35],
    )

    add_heading(doc, "十二、參考來源")
    add_bullets(
        doc,
        [
            "LINE Developers - Use rich menus: https://developers.line.biz/en/docs/messaging-api/using-rich-menus/",
            "LINE Developers - Messaging API reference / Rich menu: https://developers.line.biz/en/reference/messaging-api/#rich-menu",
            "LINE Developers - Messaging API reference / Rich menu switch action: https://developers.line.biz/en/reference/messaging-api/#richmenu-switch-action",
            "LINE Developers - LIFF overview: https://developers.line.biz/en/docs/liff/overview/",
            "專案內設計文件：docs/superpowers/specs/2026-07-03-line-rich-menu-architecture-design.md",
            "專案內實作報告：docs/superpowers/reports/2026-07-04-line-richmenu-shared-orchestrator-implementation-report.md",
        ],
    )

    return doc


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = build_doc()
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
