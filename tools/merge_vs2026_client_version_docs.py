from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


INPUT_DIR = Path(r"E:\電子書籍\改善 GitHub 多客戶版本上線追蹤")
INPUT_FILES = [
    INPUT_DIR / "VS2026_IDE_Git_Tag_Push_Handholding_Guide_2026-04-13.docx",
    INPUT_DIR / "VS2026_GitHub_Client_Version_Management_Guide_2026-04-13_FIXED.docx",
    INPUT_DIR / "VS2026_多客戶版本管理實戰手冊.docx",
    INPUT_DIR / "VS2026_GitHub_Client_Version_Management_Guide_2026-04-13.docx",
]
OUTPUT_FILE = INPUT_DIR / "VS2026_GitHub_多客戶版本管理_保母級整合教學_2026-04-13.docx"


def set_font(run, name="Microsoft JhengHei", size=None, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.bold = bold
    if size is not None:
        run.font.size = Pt(size)


def style_document(doc):
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft JhengHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        if style_name == "Normal":
            style.font.size = Pt(11)
        elif style_name == "Title":
            style.font.size = Pt(20)
        elif style_name == "Heading 1":
            style.font.size = Pt(16)
        elif style_name == "Heading 2":
            style.font.size = Pt(13)


def add_text(doc, text, style=None, bold=False, align=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_font(r, bold=bold)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(text)
    set_font(r)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_font(r)


def add_code(doc, text):
    for line in text.strip("\n").splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        set_font(r, name="Consolas", size=10)


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break()


def extract_nonempty_paragraphs(path: Path):
    doc = Document(str(path))
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]


def build_summary(paths):
    summary = []
    for path in paths:
        exists = path.exists()
        lines = extract_nonempty_paragraphs(path)[:6] if exists else []
        summary.append((path.name, exists, lines))
    return summary


def add_source_summary(doc, summary):
    doc.add_heading("附錄 A. 本次整合來源文件", level=1)
    add_text(doc, "以下 4 份文件已整合成這一本，方便你之後只看一份就能操作。")
    for name, exists, lines in summary:
        add_text(doc, f"來源：{name}", style="Heading 2")
        add_bullet(doc, f"檔案存在：{'是' if exists else '否'}")
        if lines:
            add_bullet(doc, f"內容重點：{lines[0]}")
            for line in lines[1:3]:
                add_bullet(doc, f"延伸內容：{line}")


doc = Document()
style_document(doc)

add_text(
    doc,
    "Visual Studio 2026 + GitHub 多客戶版本管理 保母級整合教學",
    style="Title",
    bold=True,
    align=WD_PARAGRAPH_ALIGNMENT.CENTER,
)
add_text(
    doc,
    "整合 4 份教學文件，濃縮成一份可以直接照做的實戰手冊",
    align=WD_PARAGRAPH_ALIGNMENT.CENTER,
)
add_text(doc, "輸出日期：2026-04-13")
add_text(doc, "適用情境：同一套系統服務 20 多個客戶，想清楚知道每個客戶目前上線哪一版，並且在 Visual Studio 2026 裡一步一步完成操作。")

doc.add_heading("1. 先講結論", level=1)
add_text(doc, "請先把這一句記住：branch 管開發線，tag 管正式上線版，clients.yaml 管每個客戶現在在線上的版本。")
add_text(doc, "你過去難管理，是因為把『客戶』『版本』『修了什麼』三件事都寫進 branch 名稱裡。以後請拆開管理。")
for item in [
    "長期 branch：client/jesus、client/sunny、client/fbllc",
    "短期 branch：feature/jesus/new-report、hotfix/jesus/fix-login",
    "正式上線 tag：deploy/jesus/5.0.9.5",
    "版本總表：deployments/clients.yaml",
]:
    add_bullet(doc, item)

doc.add_heading("2. 你最終要做到的目標", level=1)
for item in [
    "任何時候都能在 10 秒內查到某個客戶目前 production 是哪一版。",
    "每次發版都留下一個正式 tag，而不是只靠 branch 名稱猜。",
    "在 Visual Studio 2026 裡可以完成 checkout、push branch、以及搭配 Terminal 做 tag。",
    "把 20 多個客戶的版本狀態集中到同一份 clients.yaml。",
]:
    add_number(doc, item)

doc.add_heading("3. 第一次建立基礎結構", level=1)
doc.add_heading("3-1. 在 Visual Studio 2026 開啟專案", level=2)
for item in [
    "開啟 Visual Studio 2026。",
    "選『開啟專案 / 解決方案』，載入 ChurchReport.sln。",
    "等待方案、NuGet、建置狀態穩定。",
]:
    add_number(doc, item)

doc.add_heading("3-2. 把會用到的 Git 視窗全部打開", level=2)
for item in [
    "上方選單點 Git。",
    "打開 Git Repository。",
    "打開 Git Changes。",
    "上方選單點 View。",
    "打開 Terminal。",
]:
    add_number(doc, item)
add_text(doc, "之後你大部分操作只會在這三個地方切換：Git Repository、Git Changes、Terminal。")

doc.add_heading("3-3. 為每個主要客戶建立長期 branch", level=2)
add_text(doc, "先不要刪舊分支。你應該從每個客戶目前最穩定的那條舊分支，建立新的長期 branch。")
for item in [
    "在 Git Repository 視窗找到某客戶目前最穩定的舊 branch。",
    "以 Jesus 為例，可從 Jesus_5.0.9.4_BlindlySpeedUp 建立新 branch。",
    "右鍵該 branch，選 New Branch from...",
    "輸入 client/jesus。",
    "建立後切換到這條新 branch。",
]:
    add_number(doc, item)
add_text(doc, "第一批先做最常維護的 5 個客戶就好，不要一次搬完全部。")

doc.add_heading("3-4. 建立版本總表 deployments/clients.yaml", level=2)
for item in [
    "在專案根目錄建立 deployments 資料夾。",
    "新增 clients.yaml。",
    "把每個客戶目前上線的 branch、tag、commit、日期、說明寫進去。",
]:
    add_number(doc, item)
add_code(
    doc,
    """
jesus:
  branch: client/jesus
  production_tag: deploy/jesus/5.0.9.5
  commit: 7ab3c91
  deployed_at: 2026-04-13
  note: solve login timeout

sunny:
  branch: client/sunny
  production_tag: deploy/sunny/5.0.2
  commit: def5678
  deployed_at: 2026-04-02
  note: CompleteLessonList
""",
)

doc.add_heading("4. 在 Visual Studio 2026 裡做一次正式發版", level=1)
add_text(doc, "下面這一段是你最常用的實戰流程。請照順序做。")

doc.add_heading("4-1. 切到客戶長期 branch", level=2)
add_text(doc, "這一步對應 `git checkout client/jesus`。")
for item in [
    "到 Git Repository 視窗。",
    "在 branch 搜尋框輸入 client/jesus。",
    "找到後右鍵。",
    "點 Checkout。",
    "確認目前 branch 已經切換成功。",
]:
    add_number(doc, item)

doc.add_heading("4-2. 如果要開發功能或修 bug，先開短期 branch", level=2)
for item in [
    "在 client/jesus 上按右鍵。",
    "選 New Branch from...",
    "功能修改輸入 feature/jesus/new-report。",
    "修 bug 輸入 hotfix/jesus/fix-login。",
    "切到新 branch 後再開始改程式。",
]:
    add_number(doc, item)

doc.add_heading("4-3. 完成修改後 commit", level=2)
for item in [
    "打開 Git Changes。",
    "確認這次異動檔案正確。",
    "輸入清楚的 commit 訊息，例如 fix(jesus): solve login timeout。",
    "按 Commit 或 Commit All。",
]:
    add_number(doc, item)

doc.add_heading("4-4. 合回 client/jesus", level=2)
for item in [
    "切回 client/jesus。",
    "找到剛剛的 hotfix 或 feature branch。",
    "右鍵它，選 merge into current branch。",
    "確認合併完成。",
]:
    add_number(doc, item)

doc.add_heading("4-5. 建立正式上線 tag", level=2)
add_text(doc, "這一步在 Visual Studio 裡最穩定的方式，是使用內建 Terminal。")
for item in [
    "切到 Terminal。",
    "確認目前 branch 是 client/jesus。",
    "輸入正式上線 tag 指令。",
]:
    add_number(doc, item)
add_code(
    doc,
    """
git checkout client/jesus
git tag deploy/jesus/5.0.9.5
""",
)

doc.add_heading("4-6. Push branch 到 GitHub", level=2)
add_text(doc, "這一步可以用 Git Changes 圖形介面完成。")
for item in [
    "回到 Git Changes。",
    "找到右上角 Push。",
    "點下去，將 client/jesus 推上遠端。",
]:
    add_number(doc, item)
add_text(doc, "如果你習慣打指令，也可以在 Terminal 輸入：")
add_code(doc, "git push origin client/jesus")

doc.add_heading("4-7. Push tag 到 GitHub", level=2)
add_text(doc, "很多人漏掉這一步，結果 GitHub 上看不到正式版本標記。")
for item in [
    "回到 Terminal。",
    "輸入 push tag 指令。",
]:
    add_number(doc, item)
add_code(doc, "git push origin deploy/jesus/5.0.9.5")

doc.add_heading("4-8. 更新 clients.yaml", level=2)
for item in [
    "打開 deployments/clients.yaml。",
    "把 production_tag 改成 deploy/jesus/5.0.9.5。",
    "把 commit 改成此次上線 commit。",
    "把 deployed_at 改成今天日期。",
    "補上 note。",
]:
    add_number(doc, item)

doc.add_heading("5. 之後要查某個客戶現在在線上哪一版，怎麼查", level=1)
for item in [
    "第一優先直接看 deployments/clients.yaml。",
    "找到客戶名稱。",
    "直接看 production_tag。",
    "需要更細節時，再用 commit hash 去 GitHub 查。",
]:
    add_number(doc, item)
add_text(doc, "以後不要再靠記憶猜『好像是某條 branch』。標準回答要變成『Jesus 現在線上是 deploy/jesus/5.0.9.5』。")

doc.add_heading("6. 針對 20 多個客戶，建議的遷移順序", level=1)
for item in [
    "第一批先整理最常維護的 5 個客戶。",
    "建立 client/客戶名 長期 branch。",
    "補上正式上線 tag。",
    "把資料寫進 clients.yaml。",
    "等新流程穩了，再逐步納入其他客戶。",
]:
    add_number(doc, item)

doc.add_heading("7. 發版前檢查清單", level=1)
for item in [
    "我現在是不是站在正確的 client branch？",
    "這次修改是不是先在 feature/hotfix branch 完成？",
    "我有沒有 merge 回 client branch？",
    "我有沒有建立新的 deploy tag？",
    "我有沒有 push branch？",
    "我有沒有 push tag？",
    "我有沒有更新 clients.yaml？",
]:
    add_number(doc, item)

doc.add_heading("8. 常見錯誤", level=1)
for item in [
    "只 push branch，忘記 push tag。",
    "直接在 client branch 上亂改，沒有 feature/hotfix 隔離。",
    "branch 名稱同時塞客戶、版本、功能說明，導致越來越亂。",
    "沒有一份單一總表記錄每個客戶 production 狀態。",
]:
    add_bullet(doc, item)

summary = build_summary(INPUT_FILES)
add_source_summary(doc, summary)

doc.add_heading("附錄 B. 你可以直接照抄的命名模板", level=1)
add_code(
    doc,
    """
長期 branch
client/jesus
client/sunny
client/fbllc

短期功能 branch
feature/jesus/new-report
feature/sunny/member-import

短期修補 branch
hotfix/jesus/fix-login
hotfix/fbllc/fix-qrcode

正式上線 tag
deploy/jesus/5.0.9.4
deploy/jesus/5.0.9.5
deploy/sunny/5.0.2
""",
)

doc.save(str(OUTPUT_FILE))
print(OUTPUT_FILE)
