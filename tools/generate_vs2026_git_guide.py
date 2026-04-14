from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "VS2026_GitHub_Client_Version_Management_Guide_2026-04-13_FIXED.docx"


def set_run_font(run, name="Microsoft JhengHei", size=None, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.bold = bold
    if size is not None:
        run.font.size = Pt(size)


def add_text(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_run_font(r)


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(text)
    set_run_font(r)


def add_code(doc, text):
    for line in text.strip("\n").splitlines():
        p = doc.add_paragraph()
        r = p.add_run(line)
        set_run_font(r, name="Consolas", size=10)
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.space_after = Pt(0)


doc = Document()
for style_name in ["Normal", "Title", "Heading 1", "Heading 2"]:
    style = doc.styles[style_name]
    style.font.name = "Microsoft JhengHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    if style_name == "Normal":
        style.font.size = Pt(11)
    elif style_name == "Title":
        style.font.size = Pt(20)
    elif style_name == "Heading 1":
        style.font.size = Pt(16)
    elif style_name == "Heading 2":
        style.font.size = Pt(13)

title = doc.add_paragraph(style="Title")
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
r = title.add_run("Visual Studio 2026 + GitHub 多客戶版本管理保母級教學")
set_run_font(r, bold=True)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
r = subtitle.add_run("適用情境：同一套系統有 20 多個客戶，常常搞不清楚目前上線的是哪個版本")
set_run_font(r)

add_text(doc, "文件日期：2026-04-13")
add_text(doc, "閱讀方式：請照順序做，不要跳步。這份教學刻意寫得很細，目標是讓你在 Visual Studio 裡一步一步做出來。")

doc.add_heading("1. 先講結論：你要改掉什麼", level=1)
add_text(doc, "你現在最痛的地方，不是 GitHub 不好用，而是把 3 種不同的資訊都塞進 branch 名稱裡了。")
for item in ["客戶名稱", "版本號", "這次修改做了什麼"]:
    add_bullet(doc, item)
add_text(doc, "像 Jesus_5.0.9.4_BlindlySpeedUp 這種分支名稱，短時間看得懂，但客戶一多就會失控。你以後要改成：")
for item in [
    "branch 只表示開發線",
    "tag 只表示正式上線版本",
    "clients.yaml 只表示每個客戶目前 production 正在線上的是哪個 tag",
]:
    add_bullet(doc, item)

doc.add_heading("2. 你要採用的新規則", level=1)
for item in [
    "長期 branch：client/客戶名，例如 client/jesus、client/sunny",
    "短期 branch：feature/客戶名/功能、hotfix/客戶名/修補",
    "正式上線 tag：deploy/客戶名/版本號，例如 deploy/jesus/5.0.9.4",
    "上線總表檔案：deployments/clients.yaml",
]:
    add_bullet(doc, item)

doc.add_heading("3. 第一次建立基礎結構（只做一次）", level=1)
doc.add_heading("3-1. 先在 Visual Studio 開啟專案", level=2)
for item in [
    "開啟 Visual Studio 2026。",
    "選『開啟專案 / 解決方案』，打開你的 ChurchReport.sln。",
    "等專案載入完成。",
    "如果右下角還在還原套件或建置，先等它穩定。",
]:
    add_number(doc, item)

doc.add_heading("3-2. 打開 Git 視窗", level=2)
for item in [
    "在上方功能表找 Git。",
    "點開後，找到 Git Repository 或 管理分支 / Manage Branches。",
    "再把 Git Changes 視窗也打開，之後 commit 會用到。",
]:
    add_number(doc, item)

doc.add_heading("3-3. 先不要刪除舊分支", level=2)
for item in [
    "你現在看到很多像 Jesus_5.0.9.4_BlindlySpeedUp 的分支，先不要刪。",
    "把它們當成歷史資料保留。",
    "我們先建立新的規則，等新流程跑順了，再決定哪些舊分支可以封存。",
]:
    add_number(doc, item)

doc.add_heading("3-4. 為每個主要客戶建立一條長期 branch", level=2)
for item in [
    "在 Git Repository 視窗找到某個客戶目前最穩定的 branch。",
    "以 Jesus 為例，可以從 Jesus_5.0.9.4_BlindlySpeedUp 建立新 branch。",
    "在該 branch 上按右鍵，選 New Branch from...。",
    "輸入新名稱：client/jesus。",
    "勾選 checkout 新分支。",
    "對 Sunny、Fbllc、NanKan、Tycanaan 等主要客戶重複做一次。",
]:
    add_number(doc, item)

doc.add_heading("3-5. 建立 deployments 資料夾與 clients.yaml", level=2)
for item in [
    "在 Solution Explorer 裡，於專案根目錄建立資料夾 deployments。",
    "在 deployments 裡新增檔案 clients.yaml。",
    "把下面內容貼進去，再依你的客戶修改。",
]:
    add_number(doc, item)
add_code(
    doc,
    """
jesus:
  branch: client/jesus
  production_tag: deploy/jesus/5.0.9.4
  commit: abc1234
  deployed_at: 2026-04-10
  note: BlindlySpeedUp

sunny:
  branch: client/sunny
  production_tag: deploy/sunny/5.0.2
  commit: def5678
  deployed_at: 2026-04-02
  note: CompleteLessonList
""",
)

doc.add_heading("4. 第一次幫現有客戶補上正式上線 tag", level=1)
doc.add_heading("4-1. 開啟 Visual Studio 內建 Terminal", level=2)
for item in [
    "在 Visual Studio 上方找 View。",
    "打開 Terminal。",
    "確認目前路徑在你的 repo 根目錄。",
]:
    add_number(doc, item)

doc.add_heading("4-2. 切到客戶 branch", level=2)
add_number(doc, "先切到 Jesus 的長期 branch。")
add_code(doc, "git checkout client/jesus")

doc.add_heading("4-3. 對目前正式上線的 commit 打 tag", level=2)
for item in [
    "請先確認這個 branch 目前指到的 commit，就是正式上線版本。",
    "如果是，就執行下面指令。",
]:
    add_number(doc, item)
add_code(
    doc,
    """
git tag deploy/jesus/5.0.9.4
git push origin deploy/jesus/5.0.9.4
""",
)

doc.add_heading("5. 以後每次修 bug 或加功能，標準流程怎麼做", level=1)
doc.add_heading("5-1. 從客戶長期 branch 拉一條短期 branch", level=2)
for item in [
    "先 checkout 到 client/jesus。",
    "在 client/jesus 上按右鍵建立新 branch。",
    "新 branch 名稱輸入：hotfix/jesus/fix-login 或 feature/jesus/new-report。",
]:
    add_number(doc, item)

doc.add_heading("5-2. 開發與 commit", level=2)
for item in [
    "開始修改程式。",
    "打開 Git Changes 視窗。",
    "確認只有你這次要提交的檔案。",
    "寫清楚 commit 訊息，例如：fix(jesus): solve login timeout。",
]:
    add_number(doc, item)

doc.add_heading("5-3. 合回客戶長期 branch", level=2)
for item in [
    "完成測試後，切回 client/jesus。",
    "找到 hotfix/jesus/fix-login。",
    "按右鍵，選 merge into current branch。",
]:
    add_number(doc, item)

doc.add_heading("5-4. 正式上線時建立新的 tag", level=2)
add_number(doc, "假設這次上線版本是 5.0.9.5。")
add_code(
    doc,
    """
git checkout client/jesus
git tag deploy/jesus/5.0.9.5
git push origin client/jesus
git push origin deploy/jesus/5.0.9.5
""",
)

doc.add_heading("5-5. 更新 clients.yaml", level=2)
for item in [
    "打開 deployments/clients.yaml。",
    "把 production_tag 改成新的 tag。",
    "把 commit 改成這次上線的 commit hash。",
    "把 deployed_at 改成今天日期。",
]:
    add_number(doc, item)

doc.add_heading("6. 之後你要查目前哪個版本在線上，怎麼查", level=1)
for item in [
    "第一優先：打開 deployments/clients.yaml。",
    "找到客戶名稱。",
    "直接看 production_tag。這就是正式上線版。",
]:
    add_number(doc, item)

doc.add_heading("7. 一句話總結", level=1)
add_text(doc, "請記住這句就好：branch 管開發線，tag 管正式上線版，clients.yaml 管每個客戶現在在線上的版本。")

doc.save(OUTPUT)
print(OUTPUT)
