from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "VS2026_IDE_Git_Tag_Push_Handholding_Guide_2026-04-13.docx"


def set_run_font(run, name="Microsoft JhengHei", size=None, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.bold = bold
    if size is not None:
        run.font.size = Pt(size)


def add_text(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_run_font(r)


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(text)
    set_run_font(r)


def add_code(doc, text):
    for line in text.strip("\n").splitlines():
        p = doc.add_paragraph()
        r = p.add_run(line)
        set_run_font(r, name="Consolas", size=10)
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.space_after = Pt(0)


doc = Document()
for style_name in ["Normal", "Title", "Heading 1", "Heading 2"]:
    style = doc.styles[style_name]
    style.font.name = "Microsoft JhengHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    if style_name == "Normal":
        style.font.size = Pt(11)
    elif style_name == "Title":
        style.font.size = Pt(20)
    elif style_name == "Heading 1":
        style.font.size = Pt(16)
    elif style_name == "Heading 2":
        style.font.size = Pt(13)

title = doc.add_paragraph(style="Title")
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
r = title.add_run("Visual Studio 2026 IDE 操作教學")
set_run_font(r, bold=True)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
r = subtitle.add_run("主題：如何在 VS 內完成 checkout、tag、push branch、push tag")
set_run_font(r)

add_text(doc, "文件日期：2026-04-13")
add_text(doc, "適合對象：想用 Visual Studio 介面操作 Git，不想只看指令的人。")

doc.add_heading("1. 這份教學要做什麼", level=1)
add_text(doc, "你要完成的其實就是下面 4 件事：")
add_code(
    doc,
    """
git checkout client/jesus
git tag deploy/jesus/5.0.9.5
git push origin client/jesus
git push origin deploy/jesus/5.0.9.5
""",
)
add_text(doc, "在 Visual Studio 2026 裡，這 4 件事會分散在兩個地方操作：")
add_bullet(doc, "Git Repository：切 branch")
add_bullet(doc, "Terminal：建立 tag、推送 tag")
add_bullet(doc, "Git Changes 或 Sync：推送 branch")

doc.add_heading("2. 先把會用到的視窗打開", level=1)
for item in [
    "開啟 Visual Studio 2026，載入你的解決方案。",
    "上方選單點 Git。",
    "打開 Git Repository。",
    "再打開 Git Changes。",
    "上方選單點 View。",
    "打開 Terminal。",
]:
    add_number(doc, item)
add_text(doc, "只要這三個視窗都開著，後面操作會很順：Git Repository、Git Changes、Terminal。")

doc.add_heading("3. 第一步：在 VS 裡切到 client/jesus", level=1)
add_text(doc, "這一步對應的是 `git checkout client/jesus`。")
for item in [
    "看左側或底部的 Git Repository 視窗。",
    "找到 branch 清單。",
    "在搜尋框輸入 client/jesus。",
    "找到 `client/jesus` 後，對它按右鍵。",
    "點選 Checkout。",
    "等 Visual Studio 切換完成。",
]:
    add_number(doc, item)
add_text(doc, "完成後，Visual Studio 通常會在狀態列或 Git 視窗顯示目前 branch 已經變成 client/jesus。")

doc.add_heading("4. 第二步：在 VS 的 Terminal 建立 tag", level=1)
add_text(doc, "這一步對應的是 `git tag deploy/jesus/5.0.9.5`。")
for item in [
    "切到下方 Terminal 視窗。",
    "確認目前目錄是你的 repo 根目錄。",
    "輸入下面指令後按 Enter。",
]:
    add_number(doc, item)
add_code(doc, "git tag deploy/jesus/5.0.9.5")
add_text(doc, "如果沒有出現錯誤訊息，通常就表示 tag 建立成功。")

doc.add_heading("5. 第三步：把 client/jesus branch 推上 GitHub", level=1)
add_text(doc, "這一步對應的是 `git push origin client/jesus`。")
add_text(doc, "你可以用兩種方式做，我建議優先用圖形介面。")
doc.add_heading("5-1. 圖形介面做法", level=2)
for item in [
    "切到 Git Changes 視窗。",
    "看視窗右上角是否有 Push 按鈕。",
    "如果有，直接點 Push。",
    "等 Visual Studio 推送完成。",
]:
    add_number(doc, item)
doc.add_heading("5-2. Terminal 做法", level=2)
add_number(doc, "如果你比較習慣打指令，就在 Terminal 輸入：")
add_code(doc, "git push origin client/jesus")

doc.add_heading("6. 第四步：把 tag 推上 GitHub", level=1)
add_text(doc, "這一步對應的是 `git push origin deploy/jesus/5.0.9.5`。")
add_text(doc, "這一步通常直接用 Terminal 最清楚，因為很多時候 Visual Studio 圖形介面不會把 tag push 做得那麼明顯。")
for item in [
    "切回 Terminal。",
    "輸入下面指令後按 Enter。",
]:
    add_number(doc, item)
add_code(doc, "git push origin deploy/jesus/5.0.9.5")

doc.add_heading("7. 做完後怎麼確認成功", level=1)
for item in [
    "Git Repository 顯示目前 branch 是 client/jesus。",
    "Git Changes 沒有 push 錯誤。",
    "Terminal 沒有出現 push 失敗訊息。",
    "到 GitHub 網站看 Tags，可以看到 deploy/jesus/5.0.9.5。",
]:
    add_number(doc, item)

doc.add_heading("8. 你可以直接照抄的完整流程", level=1)
add_text(doc, "如果你已經在 VS 裡開好專案，實際上你只要這樣走：")
for item in [
    "Git Repository -> 找到 client/jesus -> 右鍵 -> Checkout",
    "Terminal -> 輸入 git tag deploy/jesus/5.0.9.5",
    "Git Changes -> 點 Push",
    "Terminal -> 輸入 git push origin deploy/jesus/5.0.9.5",
]:
    add_number(doc, item)

doc.add_heading("9. 常見錯誤", level=1)
for item in [
    "錯誤：還沒切到 client/jesus 就打 tag。避免方式：先看目前 branch 名稱。",
    "錯誤：只有 push branch，忘了 push tag。避免方式：最後一定再執行一次 push tag。",
    "錯誤：tag 名稱打錯。避免方式：固定用 deploy/客戶名/版本號。",
]:
    add_bullet(doc, item)

doc.add_heading("10. 一句話記住", level=1)
add_text(doc, "在 Visual Studio 裡，checkout 用 Git Repository，push branch 用 Git Changes，tag 與 push tag 用 Terminal。")

doc.save(OUTPUT)
print(OUTPUT)
